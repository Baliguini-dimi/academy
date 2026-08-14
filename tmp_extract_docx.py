from docx import Document

source = r"C:\Users\dimit\Downloads\INDUSTRIE\INDUSTRIE\INDUSTRIALISATION.docx"
document = Document(source)

for index, paragraph in enumerate(document.paragraphs, 1):
    text = paragraph.text.strip()
    if text:
        print(f"P{index}: {text}")

for table_index, table in enumerate(document.tables, 1):
    print(f"TABLE {table_index}")
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " ") for cell in row.cells))
