import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".docx_deps"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = os.path.join(os.path.dirname(__file__), "Rapport_Academy_Industrialisation_Developpement_Logiciel.docx")
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F2F2"


def set_run_font(run, size=None, bold=None, italic=None, color=BLACK):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid
    for grid_col, width in zip(tbl_grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def border_bottom(paragraph, color="000000", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, 9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.5

for style_name, size, before, after in (("Heading 1", 16, 18, 10), ("Heading 2", 14, 12, 6), ("Heading 3", 12, 8, 4)):
    style = styles[style_name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = BLACK
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Number"):
    style = styles[style_name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.25

# Cover page - restrained academic presentation.
for line, size, bold in [
    ("INSTITUT UNIVERSITAIRE D'ABIDJAN", 13, True),
    ("Master 1 - Génie Informatique / MIAGE", 12, False),
    ("Cours : Industrialisation du Développement Logiciel", 12, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(line)
    set_run_font(r, size=size, bold=bold)

doc.add_paragraph().paragraph_format.space_after = Pt(72)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
r = p.add_run("ACADEMY")
set_run_font(r, size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Industrialisation du développement logiciel appliquée à un projet de gestion académique")
set_run_font(r, size=16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(50)
r = p.add_run("Analyse du passage d'une réalisation artisanale à une démarche structurée, itérative et industrialisable")
set_run_font(r, size=12, italic=True, color=GRAY)

for label, value in [
    ("Présenté par :", "L'équipe du projet Academy"),
    ("Sous la supervision de :", "Docteur KANGA Koffi"),
    ("Année académique :", "2025-2026"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label + " ")
    set_run_font(r, 12, bold=True)
    r = p.add_run(value)
    set_run_font(r, 12)

doc.add_page_break()

# Header and footer for body pages.
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_p.paragraph_format.space_after = Pt(3)
r = header_p.add_run("Academy - Industrialisation du développement logiciel")
set_run_font(r, 9, color=GRAY)
border_bottom(header_p, color="777777", size="4")
add_page_number(section.footer.paragraphs[0])

doc.add_heading("Résumé", level=1)
add_paragraph(doc, "Ce rapport montre comment le projet Academy, une application de bureau de gestion des étudiants, des matières et des notes, s'inscrit dans le thème de l'industrialisation du développement logiciel. L'analyse part des pratiques concrètes du projet : architecture JavaFX organisée selon MVC et DAO, gestion des dépendances et du build par Maven, schéma SQLite versionné, séparation entre interface, logique applicative et accès aux données, ainsi qu'améliorations progressives des fonctionnalités et de l'ergonomie. Elle explique en quoi cette démarche réduit les dépendances aux actions individuelles, favorise la reproductibilité et prépare le projet à des contrôles qualité plus systématiques.")
add_paragraph(doc, "Le modèle retenu est itératif et incrémental, avec une conduite de projet inspirée des méthodes agiles. Ce choix est plus cohérent que le modèle en cascade ou le modèle en V, car Academy a été enrichi par livraisons successives : socle technique, gestion des étudiants, matières, notes, tableau de bord, paramètres et ajustements d'interface. Enfin, le document distingue les éléments DevOps et DevSecOps déjà préparés dans le projet de ceux qui restent à industrialiser, notamment l'intégration continue, les tests automatisés, l'analyse statique, la gestion des secrets et les sauvegardes contrôlées.")

doc.add_heading("Introduction", level=1)
add_paragraph(doc, "Le développement artisanal repose principalement sur les habitudes individuelles : chacun construit, teste et corrige selon sa propre méthode. Cette approche peut sembler rapide au démarrage, mais elle devient fragile dès que le code évolue, que plusieurs personnes interviennent ou qu'il faut reproduire une livraison. L'industrialisation répond à cette limite en installant des pratiques explicites, des outils communs et des mécanismes de contrôle qui rendent la production logicielle plus fiable et plus prévisible [1].")
add_paragraph(doc, "Le projet Academy a pour objectif de gérer les informations essentielles d'un établissement : fiches étudiantes, matières, coefficients, notes et indicateurs de suivi. Il a été développé en Java 17, avec JavaFX 21 pour l'interface, SQLite pour le stockage local et Maven pour la construction du projet. Cette combinaison constitue un cas concret permettant de relier les notions du cours à une réalisation fonctionnelle [8].")
add_paragraph(doc, "L'étude est organisée en deux parties. La première montre comment une démarche industrialisée a soutenu la réalisation du projet, en opposition à une conduite artisanale. La seconde justifie le choix du modèle itératif et incrémental à orientation agile, puis situe les perspectives DevOps et DevSecOps du projet.")

doc.add_heading("Première partie - L'industrialisation au service du projet Academy", level=1)
doc.add_heading("1. De l'approche artisanale à une réalisation structurée", level=2)
add_paragraph(doc, "Dans une approche artisanale, le développement d'Academy aurait pu consister à créer directement des fenêtres, placer les requêtes SQL dans les contrôleurs et construire le livrable depuis le poste de chaque développeur. Une telle organisation aurait rapidement produit des dépendances difficiles à maîtriser : duplication du code, erreurs lors des modifications, incohérences entre l'interface et les données, difficulté à retrouver l'origine d'un défaut et impossibilité de reconstruire le même résultat dans des conditions identiques.")
add_paragraph(doc, "Le projet a au contraire été structuré autour de composants séparés. Les vues FXML décrivent l'interface, les contrôleurs Java gèrent les interactions, les modèles représentent les données du domaine, les DAO concentrent l'accès aux données et DatabaseManager centralise l'initialisation de SQLite. Cette séparation ne constitue pas seulement un choix esthétique : elle rend chaque responsabilité plus lisible, plus testable et plus facile à faire évoluer.")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
set_table_widths(table, [2200, 3550, 3610])
headers = ["Dimension", "Développement artisanal", "Contribution de l'organisation d'Academy"]
for cell, text in zip(table.rows[0].cells, headers):
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 10, bold=True)
set_repeat_table_header(table.rows[0])
rows = [
    ("Architecture", "Logique mélangée dans les écrans.", "Séparation MVC/DAO entre interface, contrôleurs, modèles et persistance."),
    ("Construction", "Compilation dépendante d'actions manuelles.", "Maven décrit les dépendances et les objectifs de build de manière reproductible."),
    ("Données", "Création manuelle et variable de la base.", "Schéma SQL et données de démonstration versionnés ; initialisation centralisée."),
    ("Évolution", "Une correction peut entraîner des effets non maîtrisés.", "Fonctionnalités ajoutées par modules, avec contrats DAO explicites."),
    ("Traçabilité", "Difficulté à identifier les modifications.", "Code source et ressources organisés dans un dépôt Git, avec fichiers identifiables."),
]
for row_values in rows:
    cells = table.add_row().cells
    for cell, text in zip(cells, row_values):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run_font(r, 10)

doc.add_heading("2. Les fondations techniques d'une production reproductible", level=2)
add_paragraph(doc, "La reproductibilité est l'un des premiers effets recherchés par l'industrialisation. Dans Academy, le fichier pom.xml fixe les dépendances principales, notamment JavaFX, SQLite JDBC, Ikonli, OpenPDF et JUnit. Maven fournit alors une description commune du projet : un membre de l'équipe n'a pas à reconstituer manuellement les bibliothèques nécessaires. Le même code source peut être compilé selon une procédure connue, ce qui limite les écarts liés aux environnements personnels.")
add_paragraph(doc, "Le stockage des données suit la même logique. Le schéma schema.sql décrit la création des structures nécessaires et les paramètres par défaut de l'établissement. Au démarrage, DatabaseManager crée le répertoire de données et initialise la base SQLite si nécessaire. Le fichier est désormais localisé dans Documents/Academy/academy.db, ce qui sépare les données d'exécution du répertoire d'installation et rend la sauvegarde plus explicite. Les scripts de données de démonstration permettent en outre de rejouer un jeu de données cohérent pour les démonstrations et les validations fonctionnelles.")
add_paragraph(doc, "La structuration des accès aux données par interfaces DAO est également importante. Par exemple, StatistiquesDAO exprime le contrat nécessaire aux statistiques, tandis que StatistiquesDaoSqlite fournit l'implémentation liée à SQLite. La méthode d'activité récente combine les inscriptions et les notes, mais reste isolée du contrôleur de tableau de bord. Cette séparation permet de faire évoluer la persistance ou d'écrire des tests ciblés sans modifier l'interface graphique.")

doc.add_heading("3. Une qualité intégrée progressivement", level=2)
add_paragraph(doc, "L'industrialisation ne consiste pas uniquement à automatiser la compilation. Elle demande d'intégrer la qualité au fil du cycle de développement. Dans Academy, plusieurs décisions vont dans ce sens : validation des formulaires par les contrôleurs, clés et contraintes dans le schéma de données, centralisation de la connexion SQLite, gestion des erreurs lors de l'accès aux données et organisation homogène des vues. Les styles CSS réutilisables participent eux aussi à la qualité du produit : ils évitent de multiplier des règles locales et facilitent une interface cohérente.")
add_paragraph(doc, "Cependant, il convient de ne pas confondre préparation à la qualité et industrialisation totalement achevée. Le projet possède JUnit comme dépendance de test, mais une suite de tests automatisés, des seuils de couverture et une analyse statique intégrée au build devraient être généralisés. Cette limite ne remet pas en cause la démarche : elle permet plutôt de formuler une trajectoire de maturité réaliste.")

doc.add_heading("4. Collaboration, documentation et maintenabilité", level=2)
add_paragraph(doc, "Une production industrielle s'appuie sur des conventions partagées. La hiérarchie du projet Academy facilite cette compréhension : les ressources FXML sont distinctes des contrôleurs, les fichiers SQL sont regroupés dans db, la feuille theme.css centralise les choix visuels et les classes Java sont réparties selon leur rôle. Le README complète cette organisation en présentant l'objectif du logiciel, les technologies, la structure et les commandes de lancement [8].")
add_paragraph(doc, "Cette documentation réduit le temps d'appropriation du projet et évite que la connaissance reste implicite. Pour renforcer cette dimension, l'équipe peut formaliser une convention de branches, définir le contenu attendu d'une revue de code et compléter le README par un guide de contribution. Les changements seraient ainsi plus faciles à relire, à justifier et à intégrer.")

doc.add_heading("Deuxième partie - Modèle de développement retenu et perspectives DevOps", level=1)
doc.add_heading("5. Analyse des modèles envisageables", level=2)
add_paragraph(doc, "Le choix d'un modèle de développement doit être cohérent avec la nature du produit et avec le degré d'incertitude rencontré pendant sa réalisation. Academy est une application dont les fonctionnalités et l'expérience utilisateur ont été affinées progressivement. Les ajustements de l'interface, de la localisation de la base, des activités récentes du tableau de bord et des formulaires illustrent cette évolution. Un modèle strictement séquentiel aurait rendu ces retours plus coûteux.")

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
set_table_widths(table, [1800, 2300, 2600, 2660])
for cell, text in zip(table.rows[0].cells, ["Modèle", "Principe", "Intérêt pour Academy", "Limite dans ce projet"]):
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 9, bold=True)
set_repeat_table_header(table.rows[0])
model_rows = [
    ("Cascade", "Phases achevées l'une après l'autre.", "Cadre simple pour planifier un périmètre stable.", "Peu adapté aux retours fréquents sur les écrans et les priorités."),
    ("Modèle en V", "Conception et tests associés dès l'origine.", "Intéressant pour formaliser la validation.", "Reste rigide et plus approprié aux systèmes très réglementés."),
    ("Itératif et incrémental", "Produit construit par cycles ; chaque cycle ajoute un élément utilisable.", "Très adapté à l'ajout progressif des modules Academy.", "Demande une priorisation et une architecture cohérente."),
    ("Agile", "Cycles courts, retours réguliers, adaptation aux changements.", "Permet d'améliorer le produit à partir des besoins observés.", "Exige une communication continue et une discipline de suivi."),
]
for row_values in model_rows:
    cells = table.add_row().cells
    for cell, text in zip(cells, row_values):
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run_font(r, 9)

doc.add_heading("5.1 Pourquoi le modèle en cascade n'a pas été retenu", level=3)
add_paragraph(doc, "Le modèle en cascade suppose que les besoins soient définis de manière complète avant la réalisation, puis que chaque étape soit validée avant de passer à la suivante [6]. Or, dans Academy, le besoin s'est précisé au cours du travail. Les éléments déjà produits ont suscité des ajustements concrets : amélioration du comportement de la fenêtre, évolution de la présentation des tableaux, uniformisation des boutons, ajout de l'activité récente et enrichissement des paramètres de l'établissement. Ces évolutions auraient été traitées comme des retours tardifs dans un cycle en cascade, avec un risque élevé de reprise.")

doc.add_heading("5.2 Pourquoi le modèle en V n'est pas le meilleur cadre principal", level=3)
add_paragraph(doc, "Le modèle en V améliore la rigueur du modèle en cascade en associant les niveaux de conception aux niveaux de test. Il serait pertinent si Academy devait satisfaire des exigences réglementaires fortes, avec une traçabilité exhaustive entre exigences, cas de test et recette. Toutefois, l'enjeu dominant du projet est l'amélioration progressive d'une application de gestion. Le modèle en V pourrait être mobilisé ponctuellement pour formaliser les tests de fonctions sensibles, mais il ne correspond pas au rythme effectif de réalisation du projet.")

doc.add_heading("6. Le choix retenu : une démarche itérative et incrémentale, inspirée de l'agile", level=2)
add_paragraph(doc, "Le modèle qui décrit le mieux Academy est le modèle itératif et incrémental. Chaque itération produit un résultat observable et utilisable, puis crée une base plus solide pour l'itération suivante. Le projet n'est donc pas construit comme un bloc unique livré à la fin ; il progresse par incréments fonctionnels. Cette logique est cohérente avec les principes agiles, qui privilégient un logiciel fonctionnel, les interactions au sein de l'équipe et l'adaptation aux changements [2].")
add_paragraph(doc, "Dans la pratique, la démarche peut être décrite de la manière suivante :")
for item in [
    "Initialiser le socle : structure Maven, configuration JavaFX, modèle de données SQLite et architecture MVC/DAO.",
    "Construire des modules : gestion des étudiants, gestion des matières, gestion des notes et calculs associés.",
    "Rendre l'état du système visible : tableau de bord, statistiques, activités récentes et paramètres de l'établissement.",
    "Améliorer l'expérience et la cohérence : CSS partagé, boutons normalisés, tableaux plus lisibles et fenêtre adaptée à l'écran.",
    "Préparer la mise en qualité : contrats DAO, scripts SQL reproductibles, dépendance de test et pistes d'automatisation.",
]:
    add_number(doc, item)
add_paragraph(doc, "Cette conduite est proche de Scrum sans qu'il soit nécessaire d'affirmer l'application complète de toutes ses cérémonies. Une formulation exacte est donc préférable : Academy a été développé selon une approche itérative et incrémentale, pilotée par des priorités fonctionnelles et inspirée des pratiques agiles. Si l'équipe souhaite formaliser Scrum, elle peut établir un backlog, travailler par sprints courts, organiser une revue de fin d'itération et tenir une rétrospective pour améliorer le processus [5].")

doc.add_heading("7. DevOps : prolonger l'industrialisation jusqu'à la livraison", level=2)
add_paragraph(doc, "DevOps rapproche le développement et l'exploitation autour d'un flux continu : planifier, coder, construire, tester, livrer, déployer, exploiter et observer. L'objectif n'est pas seulement de livrer vite ; il est de livrer de façon répétable, mesurable et réversible [3]. Pour Academy, la présence de Maven est un premier point d'appui : la compilation, les tests et la production d'un artefact peuvent être décrits par des commandes standardisées.")
add_paragraph(doc, "Le projet est toutefois aujourd'hui plus proche d'une démarche d'industrialisation préparée que d'une chaîne DevOps complète. Une évolution progressive peut être organisée sans surdimensionner le contexte d'une application de bureau :")
for item in [
    "déclencher automatiquement mvn test et mvn package à chaque proposition de modification ;",
    "conserver l'artefact construit et vérifier qu'il démarre avec les ressources attendues ;",
    "publier une version identifiée et documenter les changements ;",
    "ajouter des tests d'intégration SQLite dans un répertoire temporaire ;",
    "mettre en place une procédure de sauvegarde et de restauration du fichier academy.db ;",
    "collecter, dans un contexte de déploiement plus large, les erreurs applicatives et les indicateurs utiles au support.",
]:
    add_bullet(doc, item)
add_paragraph(doc, "Cette trajectoire correspond à une logique de livraison continue : le logiciel reste dans un état déployable et vérifié, même si la distribution finale peut rester une décision humaine. Elle constitue une adaptation raisonnable du DevOps à un projet JavaFX local, sans imposer artificiellement des outils d'exploitation destinés à une plateforme web à grande échelle.")

doc.add_heading("8. DevSecOps : intégrer la sécurité dès la conception", level=2)
add_paragraph(doc, "DevSecOps étend DevOps en intégrant les exigences de sécurité dans le flux de développement, au lieu de les traiter uniquement avant une mise en production. Le principe est d'identifier les problèmes le plus tôt possible : dépendances vulnérables, mots de passe ou secrets exposés, validation insuffisante des entrées, fichiers de sauvegarde non protégés et erreurs de configuration [3].")
add_paragraph(doc, "Dans Academy, plusieurs points méritent une attention particulière. La base SQLite contient des données d'étudiants ; son emplacement, sa sauvegarde et les droits d'accès doivent être maîtrisés. Les formulaires doivent valider les données saisies avant persistance. Les requêtes préparées utilisées par les DAO réduisent le risque d'injection SQL. Enfin, les paramètres de l'établissement et les futures informations de connexion ne doivent jamais être placés directement dans le code source.")
add_paragraph(doc, "Les améliorations DevSecOps proposées sont les suivantes :")
for item in [
    "ajouter une analyse des dépendances Maven et corriger régulièrement les composants vulnérables ;",
    "introduire une analyse statique dans le pipeline afin de détecter des pratiques risquées ;",
    "écrire des tests de validation pour les saisies, les bornes de notes et les opérations de suppression ;",
    "chiffrer ou protéger les sauvegardes lorsqu'elles quittent le poste utilisateur ;",
    "documenter les droits nécessaires au dossier Documents/Academy et la procédure de restauration ;",
    "établir une revue de code incluant explicitement des critères de sécurité et de confidentialité.",
]:
    add_bullet(doc, item)

doc.add_heading("9. Feuille de route de maturité", level=2)
add_paragraph(doc, "La valeur de l'industrialisation se mesure dans la durée. Academy dispose déjà d'un socle utile ; l'objectif suivant est de transformer ce socle en chaîne de qualité explicite. La progression ci-dessous propose des étapes proportionnées au projet.")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
set_table_widths(table, [1900, 3900, 3560])
for cell, text in zip(table.rows[0].cells, ["Horizon", "Action prioritaire", "Résultat attendu"]):
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 10, bold=True)
set_repeat_table_header(table.rows[0])
for row_values in [
    ("Court terme", "Créer des tests unitaires pour les DAO, les calculs et les validations de formulaires.", "Régression détectée avant la démonstration ou la livraison."),
    ("Court terme", "Automatiser build et tests dans une intégration continue.", "Même vérification appliquée à chaque modification."),
    ("Moyen terme", "Versionner les livrables et formaliser la sauvegarde de la base.", "Distribution et restauration plus fiables."),
    ("Moyen terme", "Ajouter analyse de code, dépendances et revue de sécurité.", "Qualité et sécurité intégrées au flux de travail."),
    ("Long terme", "Mettre en place un suivi d'incidents et des indicateurs de qualité.", "Amélioration continue fondée sur des données réelles."),
]:
    cells = table.add_row().cells
    for cell, text in zip(cells, row_values):
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run_font(r, 10)

doc.add_heading("Conclusion", level=1)
add_paragraph(doc, "Le projet Academy montre que l'industrialisation du développement logiciel ne se limite pas aux grandes plateformes ou aux déploiements complexes. Elle commence par des décisions concrètes : organiser le code, décrire les dépendances, versionner le schéma de données, séparer les responsabilités, documenter le projet et faire évoluer le produit par incréments contrôlés. Ces choix rendent le travail plus collectif, plus reproductible et plus maintenable que dans une réalisation artisanale.")
add_paragraph(doc, "Le modèle le plus adapté au projet est l'approche itérative et incrémentale, inspirée des méthodes agiles. Elle reflète fidèlement la construction progressive des fonctions d'Academy et la prise en compte continue des retours. Le modèle en cascade et le modèle en V apportent des repères utiles de planification et de validation, mais leur rigidité ne correspond pas au déroulement réel du projet.")
add_paragraph(doc, "Enfin, DevOps et DevSecOps constituent la suite logique de cette démarche. Academy possède des fondations adaptées, notamment Maven, une architecture organisée et un schéma de base reproductible. La mise en place de tests automatisés, d'une intégration continue, d'analyses de sécurité et de procédures de sauvegarde permettra de transformer ce socle en une chaîne de développement plus mature. Le projet devient ainsi une illustration concrète du passage d'un développement artisanal à une production logicielle structurée et améliorable.")

doc.add_heading("Références bibliographiques", level=1)
references = [
    "[1] KANGA Koffi. Industrialisation du développement logiciel : chapitres 1 et 2. Support de cours, Institut Universitaire d'Abidjan, 2025-2026.",
    "[2] Beck, K. et al. Manifeste pour le développement Agile de logiciels. 2001. https://agilemanifesto.org/iso/fr/manifesto.html.",
    "[3] Kim, G., Humble, J., Debois, P. et Willis, J. The DevOps Handbook. IT Revolution Press, 2016.",
    "[4] Humble, J. et Farley, D. Continuous Delivery: Reliable Software Releases through Build, Test, and Deployment Automation. Addison-Wesley, 2010.",
    "[5] Schwaber, K. et Sutherland, J. The Scrum Guide. 2020.",
    "[6] Royce, W. W. Managing the Development of Large Software Systems. Proceedings of IEEE WESCON, 1970.",
    "[7] Forsgren, N., Humble, J. et Kim, G. Accelerate: The Science of Lean Software and DevOps. IT Revolution Press, 2018.",
    "[8] Projet Academy. README.md, pom.xml, schéma SQLite et code source du projet. Consultés dans le dépôt local, août 2026.",
]
for ref in references:
    p = add_paragraph(doc, ref)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)

doc.core_properties.title = "Academy - Industrialisation du développement logiciel"
doc.core_properties.subject = "Analyse académique du projet Academy"
doc.core_properties.author = "Équipe du projet Academy"
doc.save(OUT)
print(OUT)
